"""Create demo documents for portfolio screenshots and smoke tests."""

from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import config

SAMPLES_DIR = Path(__file__).resolve().parent.parent / "samples"


def create_docx(path: Path) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("D-GITALCODE ExtractorX — Demo Document", level=1)
    doc.add_paragraph(
        "Sample Word document with embedded brand assets for portfolio captures."
    )
    logo = config.LOGO_PNG
    if logo.exists():
        doc.add_picture(str(logo), width=Inches(2.2))
    doc.add_paragraph("Second embedded image:")
    if logo.exists():
        doc.add_picture(str(logo), width=Inches(1.4))
    doc.save(path)


def create_pdf(path: Path) -> None:
    import fitz

    logo = config.LOGO_PNG
    if not logo.exists():
        return

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_text(
        (72, 72),
        "D-GITALCODE ExtractorX — Demo PDF",
        fontsize=18,
        color=(0.15, 0.82, 0.4),
    )
    rect = fitz.Rect(72, 120, 272, 320)
    page.insert_image(rect, filename=str(logo))
    rect2 = fitz.Rect(300, 120, 500, 320)
    page.insert_image(rect2, filename=str(logo))
    doc.save(path)
    doc.close()


def main() -> None:
    SAMPLES_DIR.mkdir(parents=True, exist_ok=True)
    create_docx(SAMPLES_DIR / "demo.docx")
    create_pdf(SAMPLES_DIR / "demo.pdf")
    print(f"Samples written to {SAMPLES_DIR}")


if __name__ == "__main__":
    main()
